import pandas as pd
import numpy as np
from docx import Document
import os
import matplotlib.pyplot as plt
from docx.shared import Pt

from configparser import ConfigParser

class ConfigPa:
    ##定义配置文件读取类
    def __init__(self,filename='config.ini'):
        """
        初始化配置文件
        :param filename:配置文件名
        """
        try:
            self.filename=filename
            self.configParser = ConfigParser()
            self.configParser.read(self.filename, encoding='utf-8')
        except Exception as e:
            print(str(e))
    def get_option(self,section,option):
        """
        读取配置文件元素
        :param section:
        :param option:
        :return:
        """
        try:
            return self.configParser.get(section,option)
        except Exception as e:
            print(str(e))

class Pandas_Util:
    def __init__(self,path):
        # 文件路径 读取文件列表
        self.path = path
        self.files = os.listdir(self.path)

    def read_file(self,filename,year):
        """
        读取Excel文件
        :param filename: 文件部分字段名
        :param year:年
        :param month:月
        :return:
        """
        try:
            year = '{}年'.format(year)
            file = [x for x in self.files if filename in x and year in x ][0]
            return pd.read_excel(self.path + file)
        except Exception as e:
            print(e)

    def dataframe_setindex(self,df,index_name):
        """
        设置索引
        :param index_name: 索引名称
        :return:
        """
        df[index_name]= pd.to_datetime(df[index_name])
        df = df.set_index(index_name)
        return df

    def dataframe_reindex(self,df,index_name):
        """
        设置索引
        :param index_name:索引名称
        :return:
        """
        df.columns = df.loc[0].values
        df = df[1:]
        df[index_name] = pd.to_datetime(df[index_name])
        df = df.set_index(index_name)
        return df

    def add_unit(self,df,col_name):
        """
        添加单位
        :param df:
        :return:
        """
        df['单位'] = ''
        team = []
        for index, row in df.iterrows():
            unit0 = ''
            if type(row[col_name]) == str:
                if "溧水" in row[col_name]:
                    unit0 = "溧水"
                elif "高淳" in row[col_name]:
                    unit0 = "高淳"
                elif '六合' in row[col_name]:
                    unit0 = "六合"
                elif "浦口" in row[col_name]:
                    unit0 = "浦口"
                elif "江宁" in row[col_name]:
                    unit0 = "江宁"
                else:
                    unit0 = "城区"
            row.单位 = unit0
            team.append(unit0)
        df['单位'] = team
        return df

    def add_unit_v2(self,df):
        df['单位'] = ''
        for index, row in df.iterrows():
            #     row['部门']
            if "溧水" in row['部门']:
                row['单位'] = "溧水"
            elif "高淳" in row['部门']:
                row['单位'] = "高淳"
            else:
                if type(row['班组']) == str:
                    if '六合' in row['班组']:
                        row['单位'] = "六合"
                    elif "江宁" in row['班组']:
                        row['单位'] = "江宁"
                    elif "浦口" in row['班组']:
                        row['单位'] = "浦口"
                    else:
                        row['单位'] = "城区"
                else:
                    row['单位'] = "城区"
        return df

def get_flag(a, b):
    # 返回 上升 不变 下降
    if a > b:
        return '上升'
    elif a < b:
        return '下降'
    else:
        return '不变'

def gengerate_note():
    configPa= ConfigPa()
    current_begin_date =configPa.get_option('current', 'begin_date')
    current_end_date = configPa.get_option('current', 'end_date')
    this_year = current_end_date.split('-')[0]

    last_begin_date = configPa.get_option('last', 'begin_date')
    last_end_date = configPa.get_option('last', 'end_date')
    last_year = last_end_date.split('-')[0]

    file_path = configPa.get_option('path','file_path')
    save_path = configPa.get_option('path','save_path')
    last_year = this_year - 1
    # this_date = '%d-%02d' % (this_year, this_month)
    # round_date = '%d-%02d' % (this_year, last_month)
    # last_date = '%d-%02d' % (last_year, this_month)
    unit = {'六合营业部': '六合', '江宁营业部': '江宁', '浦口营业部': '浦口', '溧水县公司': '溧水', '配电运检室': '城区', '高淳县公司': '高淳'}

    # 2018 中压故障工单
    pandas_util= Pandas_Util(file_path)
    df_mid_vol_fault_this=pandas_util.read_file("配网故障",this_year)
    # 2017 中压故障工单
    df_mid_vol_fault_last = pandas_util.read_file("配网故障",last_year)

    df_mid_vol_fault_this = df_mid_vol_fault_this.set_index('故障时间')
    df_mid_vol_fault_last = df_mid_vol_fault_last.set_index('故障时间')
    df_mid_vol_fault_current_month = df_mid_vol_fault_this.loc[current_begin_date:current_end_date]  # 本期
    df_mid_vol_fault_round_month = df_mid_vol_fault_this[round_date]  # 上期
    df_mid_vol_fault_last_month = df_mid_vol_fault_last.loc[last_begin_date:last_end_date]  # 去年同期

    df_low_vol_fault_current = pandas_util.read_file('报修查询', this_year)
    df_low_vol_fault_current = pandas_util.dataframe_setindex(df_low_vol_fault_current, '接单登记时间')

    df_low_vol_fault_last = pandas_util.read_file('报修查询', last_year)
    df_low_vol_fault_last = pandas_util.dataframe_setindex(df_low_vol_fault_last, '接单时间')

    ## 保护动作
    dict_pro_act_current_month = {'接地': 0, '无重合闸': 0, '重合不成': 0, '重合成功': 0}
    for k, v in df_mid_vol_fault_current_month.groupby(['保护动作情况']).groups.items():
        dict_pro_act_current_month[k] = len(v)
    dict_fault_unit_current_month = {}
    dict_fault_unit_round_month = {}
    dict_fault_unit_last_month = {}
    # 本期
    for k, v in df_mid_vol_fault_current_month.groupby('单位').groups.items():
        dict_fault_unit_current_month[unit[k]] = len(v)
    # 上期
    for k, v in df_mid_vol_fault_round_month.groupby('单位').groups.items():
        dict_fault_unit_round_month[unit[k]] = len(v)
    # 去年同期
    for k, v in df_mid_vol_fault_last_month.groupby('单位').groups.items():
        dict_fault_unit_last_month[unit[k]] = len(v)
    table1 = {}
    unit_sort = ['城区', '江宁', '六合', '浦口', '溧水', '高淳', '合计']
    for unit_s in unit_sort[:-1]:
        table1[unit_s] = []
        table1[unit_s].append(dict_fault_unit_current_month[unit_s])
        table1[unit_s].append(dict_fault_unit_round_month[unit_s])
        table1[unit_s].append(dict_fault_unit_last_month[unit_s])
        table1[unit_s].append("%.2f%%" % (
                (dict_fault_unit_current_month[unit_s] - dict_fault_unit_round_month[unit_s]) /
                dict_fault_unit_round_month[unit_s] * 100))
        table1[unit_s].append("%.2f%%" % (
                (dict_fault_unit_current_month[unit_s] - dict_fault_unit_last_month[unit_s]) /
                dict_fault_unit_last_month[unit_s] * 100))
        table1[unit_s].append(
            "%.2f%%" % ((dict_fault_unit_current_month[unit_s] / dict_fault_unit_last_month[unit_s] * 100)))
    table1['合计'] = []
    table1['合计'].append(sum(dict_fault_unit_current_month.values()))
    table1['合计'].append(sum(dict_fault_unit_round_month.values()))
    table1['合计'].append(sum(dict_fault_unit_last_month.values()))
    table1['合计'].append(
        "%.2f%%" % ((sum(dict_fault_unit_current_month.values()) - sum(dict_fault_unit_round_month.values())) / sum(
            dict_fault_unit_round_month.values()) * 100))
    table1['合计'].append("%.2f%%" % (
            (sum(dict_fault_unit_current_month.values()) - sum(dict_fault_unit_last_month.values())) / sum(
        dict_fault_unit_last_month.values()) * 100))
    table1['合计'].append("%.2f%%" % (.2))
    # import time
    # name = save_path + 'report.docx'
    document = Document(docx=os.path.join(os.getcwd(), 'default.docx'))
    document.styles['Normal'].font.name = u'宋体'
    s_2_1 = "（1）中压故障数量统计\n 本期南京中压配网共发生故障{}起，其中跳闸{}起（重合成功{}起、重合不成{}起、无重合闸{}起），接地{}起。故障线路所占比例从低到高为{}。本期故障环比{}{:.2f}%，同比{}{:.2f}%。".format(
        sum(dict_pro_act_current_month.values()),
        sum(dict_pro_act_current_month.values()) - dict_pro_act_current_month['接地'],
        dict_pro_act_current_month['重合成功'], dict_pro_act_current_month['重合不成'],
        dict_pro_act_current_month['无重合闸'], dict_pro_act_current_month['接地'],
        '、'.join([x[0] for x in sorted(dict_fault_unit_current_month.items(), key=lambda a: a[1])]),
        get_flag(sum(dict_fault_unit_current_month.values()), sum(dict_fault_unit_round_month.values())),
        abs(sum(dict_fault_unit_current_month.values()) - sum(dict_fault_unit_round_month.values())) / sum(
            dict_fault_unit_round_month.values()) * 100,
        get_flag(sum(dict_fault_unit_current_month.values()), sum(dict_fault_unit_last_month.values())),
        abs(sum(dict_fault_unit_current_month.values()) - sum(dict_fault_unit_last_month.values())) / sum(
            dict_fault_unit_last_month.values()) * 100)
    document.add_paragraph(s_2_1)

    table = document.add_table(rows=9, cols=7, style='Table Grid')
    table.style.font.size = Pt(12)

    table.cell(0, 0).merge(table.cell(0, 6))
    table.rows[0].cells[0].text = '表1 中压故障对比数据'
    con = ['单位', '本期数据', '上期数据', '去年同期', '环比', '同比', '线路比例']
    for row, obj_row in enumerate(table.rows[1:2]):
        for col, cell in enumerate(obj_row.cells):
            cell.text = con[col]
    for row, obj_row in enumerate(table.rows[2:9]):
        for col, cell in enumerate(obj_row.cells[0:1]):
            cell.text = unit_sort[row]
    for row, obj_row in enumerate(table.rows[2:]):
        for col, cell in enumerate(obj_row.cells[1:]):
            cell.text = str(table1[unit_sort[row]][col])

    # 中压故障设备统计
    df_fault_unknown = df_mid_vol_fault_current_month['涉及设备'].dropna()
    dict_fault_equ = dict([(x, list(df_fault_unknown).count(x)) for x in list(df_fault_unknown)])
    sorted_fault_equ = sorted(dict_fault_equ.items(), key=lambda a: a[1], reverse=True)
    dict_equ_fault_unit = {}
    for key, value in unit.items():
        dict_equ_fault_unit[value] = {}
    for k, v in df_mid_vol_fault_current_month.groupby(['单位', '涉及设备']).groups.items():
        if type(k[1]) == str:
            dict_equ_fault_unit[unit[k[0]]][k[1]] = len(v)
    s_2_2 = "（2）中压故障设备统计\n中压故障涉及设备{}个，另有{}个无法明确故障设备。其中{}、{}占比较大，共占故障总设备数的{:2.0f}%。各区域主要故障设备分别是{}。"
    s_part = ""
    plt.rcParams['font.sans-serif'] = ['SimHei']
    p_row = 321
    d_fault_equ = {}
    for u in unit_sort[:-1]:
        d_fault_equ[u] = []
    dd_fault_equ = {}
    for key, value in dict_equ_fault_unit.items():
        v = sorted(value.items(), key=lambda a: a[1], reverse=True)
        labels = value.keys()
        fracs = value.values()
        d_fault_equ[key].append(fracs)
        d_fault_equ[key].append(labels)
        for k, vv in value.items():
            if k not in dd_fault_equ:
                dd_fault_equ[k] = 0
            dd_fault_equ[k] += vv
        if sum(value.values()) <= 3:
            s_part += key
            for vv in v:
                s_part += vv[0] + str(vv[1]) + '起，'
        else:
            s_part += key + '(共%d起)' % (sum(value.values())) + v[0][0] + '占%2.0f%%' % (
                    v[0][1] / sum(value.values()) * 100) + ','
    s_2_2 = s_2_2.format(len(df_mid_vol_fault_current_month),
                         len(df_mid_vol_fault_current_month) - len(df_fault_unknown),
                         sorted_fault_equ[0][0], sorted_fault_equ[1][0],
                         (sorted_fault_equ[0][1] + sorted_fault_equ[1][1]) / len(df_fault_unknown) * 100, s_part[:-1])
    plt.pie(x=dd_fault_equ.values(), autopct='%2.0f%%',shadow=False, labeldistance=5, startangle=90, pctdistance=1.1,center=(-25,-20))
    plt.legend(dd_fault_equ.keys(), loc='lower right', bbox_to_anchor=(1.35, 0.3))
    plt.title('全市', fontsize=20)
    fig = plt.gcf()
    fig.savefig("all_equ_fault.png")
    plt.figure(figsize=(10, 10))
    for u in unit_sort[:-1]:
        plt.subplot(p_row)
        fracs = d_fault_equ[u][0]
        labels = d_fault_equ[u][1]
        plt.pie(x=fracs, autopct='%2.0f%%',shadow=False, labeldistance=5, startangle=90, pctdistance=1.15,center=(-10,-10))
        plt.legend(labels, loc='lower right', bbox_to_anchor=(1.6, 0.3))
        plt.title(u, fontsize=20, loc='center')
        p_row += 1
    fig = plt.gcf()
    fig.savefig("each_equ_fault.png")
    plt.close('all')
    document.add_picture('all_equ_fault.png')
    document.add_picture('each_equ_fault.png')
    document.add_paragraph(s_2_2)
    # 中压故障原因统计
    dict_equ_fault_weather_unit = {}
    dict_equ_weather = {}
    unknown_cause = 0

    for key, value in unit.items():
        dict_equ_fault_weather_unit[value] = {}
    for k, v in df_mid_vol_fault_current_month.groupby(['单位', '故障分类']).groups.items():
        k = list(k)
        if type(k[1]) == str:
            if k[1] == '原因不明': unknown_cause += len(v)
            elif k[1] not in dict_equ_fault_weather_unit[unit[k[0]]]:
                dict_equ_fault_weather_unit[unit[k[0]]][k[1]] = 0
                dict_equ_fault_weather_unit[unit[k[0]]][k[1]] += len(v)
                if k[1] not in dict_equ_weather: dict_equ_weather[k[1]] = 0
                dict_equ_weather[k[1]] += len(v)
    s = sorted(dict_equ_weather.items(), key=lambda a: a[1], reverse=True)
    s_2_3_part = ""
    for key, value in dict_equ_fault_weather_unit.items():
        v = sorted(value.items(), key=lambda a: a[1], reverse=True)
        if sum(value.values()) <= 3:
            s_2_3_part += key
            for vv in v:
                s_2_3_part += vv[0] + str(vv[1]) + '起，'
        else:
            s_2_3_part += key + '(共%d起)' % (sum(value.values())) + v[0][0] + '占%2.0f%%' % (
                    v[0][1] / sum(value.values()) * 100) + ','
    cause_value = dict_equ_weather.values()
    cause_key=dict_equ_weather.keys()
    plt.pie(x=cause_value, autopct='%2.0f%%',shadow=False, labeldistance=5, startangle=90, pctdistance=1.1,center=(25,20))
    plt.legend(cause_key, loc='lower right', bbox_to_anchor=(1.35, 0.3))
    plt.title('全市', fontsize=20)
    # plt.show()
    fig = plt.gcf()
    fig.savefig("all_equ_fault_cause.png", dpi=300)

    plt.figure(figsize=(10, 10))
    p_row = 321
    for u in unit_sort[:-1]:
        plt.subplot(p_row)
        fracs = dict_equ_fault_weather_unit[u].values()
        labels = dict_equ_fault_weather_unit[u].keys()
        plt.pie(x=fracs, autopct='%2.0f%%',
                shadow=False, labeldistance=5, startangle=90, pctdistance=1.15
                )
        plt.legend(labels, loc='lower right', bbox_to_anchor=(1.6, 0.3))
        plt.title(u, fontsize=20, loc='center')
        p_row += 1

    fig = plt.gcf()
    fig.savefig("each_equ_fault_cause.png", dpi=300)
    plt.close('all')
    document.add_picture('all_equ_fault_cause.png')
    document.add_picture('each_equ_fault_cause.png')

    s_2_3 = "（3）中压故障原因统计\n在各中压故障原因中有{}起故障无法明确原因，其中{}、{}、{}占比较大，各占中压故障数量的{}、{}、{}。各区域主要故障原因分别是{}。"
    s_2_3 = s_2_3.format(unknown_cause, s[0][0], s[1][0], s[2][0],
                         '%2.0f%%' % ((s[0][1]) / (sum(dict_equ_weather.values())) * 100),
                         '%2.0f%%' % ((s[1][1]) / (sum(dict_equ_weather.values())) * 100),
                         '%2.0f%%' % ((s[2][1]) / (sum(dict_equ_weather.values())) * 100)
                         , s_2_3_part[:-1])
    document.add_paragraph(s_2_3)

    ##（4）低压故障数量统计
    ####  2018 加入单位
    df_low_vol_fault_current = pandas_util.add_unit(df_low_vol_fault_current,'抢修队伍/班组')
    ##3 2017 加入单位
    df_low_vol_fault_last = pandas_util.add_unit(df_low_vol_fault_last,'供电单位')
    ### 低压 本期
    df_low_vol_fault_current_month = df_low_vol_fault_current.loc[current_begin_date:current_end_date]
    ### 低压上期
    df_low_vol_fault_round_month = df_low_vol_fault_current[round_date]
    #### 低压 去年同期
    df_low_vol_fault_last_month = df_low_vol_fault_last.loc[last_begin_date:last_end_date]
    df_tmp_current = df_low_vol_fault_current_month.where(df_low_vol_fault_current_month['关联主单'].isnull()).where(
        ~df_low_vol_fault_current_month['记录到达时间'].isnull()).where(~df_low_vol_fault_current_month['通知抢修时间'].isnull())
    df_tmp_current = df_tmp_current.loc[df_tmp_current['一级分类'] == '低压故障']
    df_tmp_round = df_low_vol_fault_round_month.where(df_low_vol_fault_round_month['关联主单'].isnull()).where(
        ~df_low_vol_fault_round_month['记录到达时间'].isnull())
    df_tmp_round = df_tmp_round.loc[df_tmp_round['一级分类'] == '低压故障']
    df_tmp_last = df_low_vol_fault_last_month.where(~df_low_vol_fault_last_month['修复时间'].isnull())
    df_tmp_last = df_tmp_last.loc[df_tmp_last['一级分类'] == '低压故障']
    dict_tmp_current = {}
    for k, v in df_tmp_current.groupby(['单位']).groups.items():
        dict_tmp_current[k] = len(v)
    dict_tmp_round = {}
    for k, v in df_tmp_round.groupby(['单位']).groups.items():
        dict_tmp_round[k] = len(v)
    dict_tmp_last = {}
    for k, v in df_tmp_last.groupby(['单位']).groups.items():
        dict_tmp_last[k] = len(v)
    table2 = {}
    for unit_s in unit_sort[:-1]:
        table2[unit_s] = []
        table2[unit_s].append(dict_tmp_current[unit_s])
        table2[unit_s].append(dict_tmp_round[unit_s])
        table2[unit_s].append(dict_tmp_last[unit_s])
        table2[unit_s].append(
            "%.2f%%" % ((dict_tmp_current[unit_s] - dict_tmp_round[unit_s]) / dict_tmp_round[unit_s] * 100))
        table2[unit_s].append(
            "%.2f%%" % ((dict_tmp_current[unit_s] - dict_tmp_last[unit_s]) / dict_tmp_last[unit_s] * 100))
    table2['合计'] = []
    table2['合计'].append(sum(dict_tmp_current.values()))
    table2['合计'].append(sum(dict_tmp_round.values()))
    table2['合计'].append(sum(dict_tmp_last.values()))
    table2['合计'].append("%.2f%%" % (
            (sum(dict_tmp_current.values()) - sum(dict_tmp_round.values())) / sum(dict_tmp_round.values()) * 100))
    table2['合计'].append(
        "%.2f%%" % ((sum(dict_tmp_current.values()) - sum(dict_tmp_last.values())) / sum(dict_tmp_last.values()) * 100))
    document.add_paragraph('\n')
    document.styles['Normal'].font.name = u'宋体'
    table = document.add_table(rows=9, cols=6, style='Table Grid')
    table.style.font.size = Pt(12)
    table.cell(0, 0).merge(table.cell(0, 5))
    table.rows[0].cells[0].text = '表2 低压故障对比数据'
    con = ['单位', '本期数据', '上期数据', '去年同期', '环比', '同比']
    for row, obj_row in enumerate(table.rows[1:2]):
        for col, cell in enumerate(obj_row.cells):
            cell.text = con[col]
    for row, obj_row in enumerate(table.rows[2:9]):
        for col, cell in enumerate(obj_row.cells[0:1]):
            cell.text = unit_sort[row]
    for row, obj_row in enumerate(table.rows[2:]):
        for col, cell in enumerate(obj_row.cells[1:]):
            cell.text = str(table2[unit_sort[row]][col])
    c, r, l = len(df_tmp_current), len(df_tmp_round), len(df_tmp_last)
    up, down, equal = [], [], []
    for k in unit_sort:
        flag = get_flag(len(df_tmp_current.loc[df_tmp_current.单位 == k]), len(df_tmp_round.loc[df_tmp_round.单位 == k]))
        if flag == '上升':
            up.append(k)
        elif flag == '下降':
            down.append(k)
        else:
            equal.append(k)
    s_1_4_1 = "同期"
    if up != []:
        s_1_4_1 += '上升的有'
        for u in up: s_1_4_1 += str(u) + ','
    if down != []:
        s_1_4_1 += '下降的有'
        for d in down: s_1_4_1 += str(d) + ','
    s_1_4 = "本期配网低压共发生故障{}起，环比{}{:.2f}%，同比{}{:.2f}%。{}。".format(
        c, get_flag(c, r), abs(c - r) / r, get_flag(c, l), abs(c - l) / l, s_1_4_1[:-1])

    document.add_paragraph(s_1_4)

    #####（5）低压故障设备及原因统计
    dd = {}
    ss = 0
    ddd = {}
    sss = 0
    for k, v in df_tmp_current.groupby(['故障原因']).groups.items():
        ddd[k] = len(v)
        sss += len(v)
    for k, v in df_tmp_current.groupby(['三级分类']).groups.items():
        dd[k] = len(v)
        ss += len(v)
    ddd_sort = sorted(ddd.items(), key=lambda a: a[1], reverse=True)
    s = 0
    dddd = {}
    for u in unit_sort[:-1]:
        dddd[u] = {}
    for k, v in df_tmp_current.groupby(['单位', '三级分类']).groups.items():
        site = k[0]
        dddd[site][k[1]] = len(v)
        s += len(v)
    s_1_5 = '本期低压故障涉及设备{}个，其中{}占比较大，共占故障设备总量的{:.1f}%。各区域低压故障设备最多的分别为'.format(len(df_tmp_current),','.join([x[0] for x in sorted(dd.items(),key=lambda a: a[1],reverse=True)[:4]]),
         sum([x[1] for x in sorted(dd.items(), key=lambda a: a[1],reverse=True)[:4]]) / len(df_tmp_current) * 100)
    for u in unit_sort[:-1]:
        s_d = sorted(dddd[u].items(), key=lambda a: a[1], reverse=True)[0]
        s_1_5 += u + s_d[0] + '占' + str(round(s_d[1] / sum(dddd[u].values()) * 100)) + '%,'
    s_1_5 += '故障原因以{}为主，占低压故障数量的{:.0f}%。'.format(ddd_sort[0][0], ddd_sort[0][1] / len(df_tmp_current) * 100)
    s_dd = sorted(dd.items(),key=lambda a:a[1],reverse=True)
    dd_vs=[]
    dd_key=[]
    if len(dd.keys())>9:
        for s_d in s_dd[:9]:
            dd_vs.append(s_d[1])
            dd_key.append(s_d[0])
        dd_vs.append(sum([i[1] for i in s_dd[9:]]))
        dd_key.append('其他')
    else:
        dd_vs=dd.values()
        dd_key = dd.keys()
    plt.figure(figsize=(7, 4))
    plt.pie(x=dd_vs, autopct='%2.0f%%', shadow=False, labeldistance=5, startangle=90, pctdistance=1.05)
    plt.legend(dd_key, loc='lower right', bbox_to_anchor=(1.57, 0.2))
    plt.title('全市', fontsize=20)
    fig = plt.gcf()
    fig.savefig("all_equ_fault_cause_low.png", dpi=300)
    plt.figure(figsize=(14, 10))
    p_row = 321
    for u in unit_sort[:-1]:
        plt.subplot(p_row)
        fracs = dddd[u].values()
        labels = dddd[u].keys()
        plt.pie(x=fracs, autopct='%2.0f%%',
                shadow=False, labeldistance=4.5, startangle=90, pctdistance=1.10
                )
        plt.legend(labels, loc='lower right', bbox_to_anchor=(1.8, 0.10))
        plt.title(u, fontsize=20, loc='center')
        p_row += 1

    fig = plt.gcf()
    fig.savefig("each_equ_fault_cause_low.png", dpi=300)
    plt.close('all')
    document.add_picture('all_equ_fault_cause_low.png')
    document.add_picture('each_equ_fault_cause_low.png')

    document.add_paragraph(s_1_5)

    ####### 指标监控################
    ####公变出口低电压######
    document.add_paragraph("4、指标监测情况")
    document.add_paragraph("（1）电压监测情况")
    document.add_paragraph("3)公变出口低电压")

    df_export_low_vol = pandas_util.read_file("公变出口", this_year)
    df_export_low_vol = pandas_util.dataframe_reindex(df_export_low_vol,'发生时刻')

    #  设置相应单位
    df_export_low_vol = pandas_util.add_unit_v2(df_export_low_vol)
    s = 0
    d_4_1_3 = {}
    for u in unit_sort[:]:
        d_4_1_3[u] = []
    for k, v in df_export_low_vol[this_date].groupby(['单位']).groups.items():
        d_4_1_3[k].append(len(v))
        d_4_1_3[k].append('{:.2f}%'.format(len(v) / float(configPa.get_option('配变',k))))
        s += len(v)
    change_sum = sum([int(configPa.get_option('配变',k)) for k in unit_sort[:-1]])
    d_4_1_3_s = sorted(d_4_1_3.items(), key=lambda a: a[1], reverse=True)
    s_4_1_3 = "本期公变出口低电压共{}台，低电压比例{:.1f}%，其中公变低电压户数最多的为{}{}台。".format(s, s / change_sum, d_4_1_3_s[0][0], d_4_1_3_s[0][1][0])
    d_4_1_3['合计'].append(s)
    d_4_1_3['合计'].append('{:.2f}%'.format(s / 10))

    table = document.add_table(rows=9, cols=3, style='Table Grid')
    table.style.font.size = Pt(12)
    table.cell(0, 0).merge(table.cell(0, 2))
    table.rows[0].cells[0].text = '表5 公变出口低电压情况'
    table.rows[1].cells[0].text = '单位'
    table.rows[1].cells[1].text = '公变低电压数'
    table.rows[1].cells[2].text = '低电压比例(%)'

    for row, obj_row in enumerate(table.rows[2:]):
        for col, cell in enumerate(obj_row.cells[0:1]):
            cell.text = unit_sort[row]
    for row, obj_row in enumerate(table.rows[2:]):
        for col, cell in enumerate(obj_row.cells[1:3]):
            if d_4_1_3[unit_sort[row]] == []:
                d_4_1_3[unit_sort[row]] = [0, '0.0%']
            cell.text = str(d_4_1_3[unit_sort[row]][col])
    document.add_paragraph(s_4_1_3)

    #########4）D类用户出口低电压############
    df_class_d_user= pandas_util.read_file("D类用户", this_year, this_month)
    df_class_d_user= pandas_util.dataframe_reindex(df_class_d_user, '发生日期')
    #  设置相应单位
    df_class_d_user  = pandas_util.add_unit_v2(df_class_d_user)
    s = 0
    document.add_paragraph('4）D类用户出口低电压')
    d_4_1_4 = {}
    for u in unit_sort[:]:
        d_4_1_4[u] = []
    for k, v in df_class_d_user[this_date].groupby(['单位']).groups.items():
        d_4_1_4[k].append(len(v))
        d_4_1_4[k].append('{:.2f}%'.format(len(v) / float(configPa.get_option('配变',k))))
        s += len(v)
    d_4_1_4_s = sorted(d_4_1_4.items(), key=lambda a: a[1], reverse=True)
    s_4_1_4 = "本期D类用户出口低电压共{}户，低电压比例{:.2f}%，其中用户低电压户数最多的为{}{}户。".format(s, s / 1000, d_4_1_4_s[0][0],
                                                                        d_4_1_4_s[0][1][0])
    d_4_1_4['合计'].append(s)
    d_4_1_4['合计'].append('{:.2f}%'.format(s / 10))
    table = document.add_table(rows=9, cols=3, style='Table Grid')
    table.style.font.size = Pt(12)
    table.cell(0, 0).merge(table.cell(0, 2))
    table.rows[0].cells[0].text = '表6 D类用户出口低电压情况'
    table.rows[1].cells[0].text = '单位'
    table.rows[1].cells[1].text = '用户低电压数'
    table.rows[1].cells[2].text = '低电压比例(%)'
    for row, obj_row in enumerate(table.rows[2:]):
        for col, cell in enumerate(obj_row.cells[0:1]):
            cell.text = unit_sort[row]
    for row, obj_row in enumerate(table.rows[2:]):
        for col, cell in enumerate(obj_row.cells[1:3]):
            if d_4_1_4[unit_sort[row]] == []:
                d_4_1_4[unit_sort[row]] = [0, '0.0%']
            cell.text = str(d_4_1_4[unit_sort[row]][col])
    document.add_paragraph(s_4_1_4)

    ###################（2）超重载配变（2）超重载配变
    document.add_paragraph('（2）超重载配变')
    df_load_change = pandas_util.read_file("配变负载", this_year)
    df_load_change = pandas_util.dataframe_reindex(df_load_change, '发生时刻')

    #  设置相应单位
    df_load_change   = pandas_util.add_unit_v2(df_load_change)
    s = 0
    d_4_2 = {}
    for u in unit_sort[:]:
        d_4_2[u] = []
    for k, v in df_load_change[this_date].groupby(['单位']).groups.items():
        d_4_2[k].append(len(v))
        d_4_2[k].append('{:.2f}%'.format(len(v) / 10))
        s += len(v)
    d_4_2_s = sorted(d_4_2.items(), key=lambda a: a[1], reverse=True)
    s_4_2 = "本期配变超重载共{}台，超重载比例{:.2f}%，其中{}最多为{}台。".format(s, s / 1000, d_4_2_s[0][0], d_4_2_s[0][1][0])
    d_4_2['合计'].append(s)
    d_4_2['合计'].append('{:.2f}%'.format(s / 10))
    table = document.add_table(rows=9, cols=3, style='Table Grid')
    table.style.font.size = Pt(12)
    table.cell(0, 0).merge(table.cell(0, 2))
    table.rows[0].cells[0].text = '表7 配变超重载情况'
    table.rows[1].cells[0].text = '单位'
    table.rows[1].cells[1].text = '配变超重载数'
    table.rows[1].cells[2].text = '配变超重载比例(%)'
    for row, obj_row in enumerate(table.rows[2:]):
        for col, cell in enumerate(obj_row.cells[0:1]):
            cell.text = unit_sort[row]
    for row, obj_row in enumerate(table.rows[2:]):
        for col, cell in enumerate(obj_row.cells[1:3]):
            if d_4_2[unit_sort[row]] == []:
                d_4_2[unit_sort[row]] = [0, '0.0%']
            cell.text = str(d_4_2[unit_sort[row]][col])
    document.add_paragraph(s_4_2)

    ########  ########### ###########（3）配变停运情况
    document.add_paragraph('（3）配变停运情况')

    df_change_stop = pandas_util.read_file("公变停运", this_year)
    df_change_stop = pandas_util.dataframe_reindex(df_change_stop, '失电时间')
    line_sum = sum([int(configPa.get_option('线路总数', k)) for k in ['城区', '江宁', '六合', '浦口', '溧水', '高淳']])
    d_stop = {}
    dd_stop = {}  ## 重复
    ddd_stop = {} ## 总共
    line_count = {}
    for k in unit_sort[:-1]:
        d_stop[k] = []
        line_count[k] = int(configPa.get_option('线路总数', k))
        dd_stop[k] = 0
        ddd_stop[k] = 0
    #  设置相应单位
    df_change_stop = pandas_util.add_unit_v2(df_change_stop)

    one = two = three = four = 0
    for k, v in df_change_stop.groupby(['单位', '重复停运次数']).groups.items():
        if int(k[1]) == 1:
            d_stop[k[0]].append(len(v))
            one += len(v)
        elif int(k[1]) == 2:
            d_stop[k[0]].append(len(v))
            two += len(v)
        elif int(k[1]) == 3:
            d_stop[k[0]].append(len(v))
            three += len(v)
        elif int(k[1]) >= 4:
            dd_stop[k[0]] += len(v)
            four += len(v)
        ddd_stop[k[0]] += len(v) * int(k[1])
    highest = ''
    max_stop = 0.0
    lowest = ''
    min_stop = 100.0
    for k in unit_sort[:-1]:
        d_stop[k].insert(0, ddd_stop[k])
        d_stop[k].append(dd_stop[k])
        d_stop[k].append(str(round(ddd_stop[k] / line_count[k] * 100, 2)))
        d_stop[k].append(str(round(dd_stop[k] / line_count[k] * 100, 2)))
        if ddd_stop[k] / line_count[k] >= max_stop:
            max_stop = ddd_stop[k] / line_count[k]
            highest = k
        if ddd_stop[k] / line_count[k] <= min_stop:
            min_stop = ddd_stop[k] / line_count[k]
            lowest = k
    d_stop['合计'] = []
    d_stop['合计'].append(sum(ddd_stop.values()))
    d_stop['合计'].append(one)
    d_stop['合计'].append(two)
    d_stop['合计'].append(three)
    d_stop['合计'].append(four)
    d_stop['合计'].append(round(sum(ddd_stop.values()) / line_sum* 100,2))
    d_stop['合计'].append(round(sum(dd_stop.values()) / line_sum* 100, 2))

    s_4_3 = "对本月南京地区配变停运情况进行统计，配变停运总次数{}台次，停运2次{}台，停运3次{}台，停运4次{}台。配变重复停运占比最高为{}{:.2f}%，最低为{}{:.2f}%。".format(sum(ddd_stop.values()), two, three, four, highest, max_stop, lowest, min_stop)
    table = document.add_table(rows=9, cols=8, style='Table Grid')
    table.style.font.size = Pt(12)
    table.cell(0, 0).merge(table.cell(0, 7))
    table.rows[0].cells[0].text = '表8 配变重复停运情况'
    table.rows[1].cells[0].text = '单位'
    table.rows[1].cells[1].text = '配变停运总次数'
    table.rows[1].cells[2].text = '停运1次'
    table.rows[1].cells[3].text = '停运2次'
    table.rows[1].cells[4].text = '停运3次'
    table.rows[1].cells[5].text = '停运4次及以上'
    table.rows[1].cells[6].text = '配变停运占比（%）'
    table.rows[1].cells[7].text = '配变重复停运占比（%）'
    for row, obj_row in enumerate(table.rows[2:]):
        for col, cell in enumerate(obj_row.cells[0:1]):
            cell.text = unit_sort[row]
    for row, obj_row in enumerate(table.rows[2:]):
        for col, cell in enumerate(obj_row.cells[1:]):
            cell.text = str(d_stop[unit_sort[row]][col])
    document.add_paragraph(s_4_3)

    ##################
    document.add_paragraph('（4）线路重复故障情况')
    df_sofar = df_mid_vol_fault_this[(df_mid_vol_fault_this.index.get_level_values(0) <= '2018-04-12')]
    line_dict = {}
    line_dict_table={}
    line_two = line_three = line_four = 0
    line_max_rate = 0.0
    line_min_rate = 1.0
    line_max_city = line_min_city = ''
    for k in ['城区', '江宁', '六合', '浦口', '溧水', '高淳']:
        line_dict[k] = {}
        line_dict_table[k] = []
        line_dict[k][2] = 0
        line_dict[k][3] = 0
        line_dict[k][4] = 0
    for k, v in df_sofar.groupby(['单位', '线路名称']).groups.items():
        k = unit[k[0]]
        if len(v) == 2:
            line_dict[k][2] += 1
            line_two += 1
        elif len(v) == 3:
            line_dict[k][3] += 1
            line_three += 1
        elif len(v) >= 4:
            line_dict[k][4] += 1
            line_four += 1
    for k in unit_sort[:-1]:
        line_dict_table[k].append(sum(line_dict[k].values()))
        rate = round(sum(line_dict[k].values()) / int(configPa.get_option('线路总数', k)) * 100, 2)
        line_dict_table[k].append(rate)
        if rate >= line_max_rate:
            line_max_rate = rate
            line_max_city = k
        if rate <= line_min_rate:
            line_min_rate = rate
            line_min_city = k
        line_dict_table[k].append(line_dict[k][2])
        line_dict_table[k].append(line_dict[k][3])
        line_dict_table[k].append(line_dict[k][4])
    line_dict_table['合计'] = []
    line_dict_table['合计'].append(sum([sum(v.values()) for v in line_dict.values()]))

    repate_rate= round(sum([sum(v.values()) for v in line_dict.values()]) * 100 / line_sum, 2)
    line_dict_table['合计'].append(repate_rate)
    line_dict_table['合计'].append(line_two)
    line_dict_table['合计'].append(line_three)
    line_dict_table['合计'].append(line_four)

    s_4_4 = "截止至{}，{}年全市线路停运2次{}条，停运3次{}条，停运4次{}条。重复停运比例{}%，其中{}最高为{}%，{}最低为{}%。".format("04-12",this_year, line_two,
                                                                                         line_three, line_four,repate_rate,
                                                                                         line_max_city,
                                                                                         line_max_rate,line_min_city,line_min_rate)
    table = document.add_table(rows=10, cols=6, style='Table Grid')
    table.style.font.size = Pt(12)
    table.cell(0, 0).merge(table.cell(0, 5))
    table.rows[0].cells[0].text = '表9 重复故障线路情况'
    table.cell(1, 0).merge(table.cell(2, 0))
    table.rows[1].cells[0].text = '单位'
    table.cell(1, 1).merge(table.cell(2, 1))
    table.rows[1].cells[1].text = '重复故障线路总数'

    table.cell(1, 2).merge(table.cell(2, 2))
    table.rows[1].cells[2].text = '重复故障线路占比（%）'
    table.cell(1, 3).merge(table.cell(1, 5))
    table.rows[1].cells[3].text = '重复情况'
    table.rows[2].cells[3].text = '2次'
    table.rows[2].cells[4].text = '3次'
    table.rows[2].cells[5].text = '4次'
    for row, obj_row in enumerate(table.rows[3:]):
        for col, cell in enumerate(obj_row.cells[0:1]):
            cell.text = unit_sort[row]
    for row, obj_row in enumerate(table.rows[3:]):
        for col, cell in enumerate(obj_row.cells[1:]):
            cell.text = str(line_dict_table[unit_sort[row]][col])
    document.add_paragraph(s_4_4)
#################
    ### 配网中压
    fault_order = {}
    for k in unit_sort:
        fault_order[k] = []
    # 各单位工单数
    n = 0
    for k, v in df_low_vol_fault_current_month.groupby(['单位']).groups.items():
        if k != '':
            fault_order[k].append(len(v))
        n += len(v)
    fault_order['合计'].append(n)
    n = 0
    ### 配网中压 各单位 故障数
    for k, v in df_mid_vol_fault_current_month.groupby(['单位']).groups.items():
        if k != '':
            fault_order[unit[k]].append(len(v))
        n += len(v)
    fault_order['合计'].append(n)
    ### 配网中压 各单位 工单数
    n = 0
    for k, v in df_low_vol_fault_current_month.loc[df_low_vol_fault_current_month.一级分类 == '高压故障'].groupby(
            ['单位']).groups.items():
        fault_order[k].append(len(v))
        n += len(v)
    fault_order['合计'].append(n)
    for k in unit_sort:
        fault_order[k].append(round(fault_order[k][2] / fault_order[k][1], 1))
        fault_order[k].append(round(fault_order[k][2] / fault_order[k][0] * 100, 2))
    ### 配网低压故障数
    n = 0
    for k, v in df_tmp_current.groupby(['单位']).groups.items():
        fault_order[k].append(len(v))
        n += len(v)
    fault_order['合计'].append(n)
    ### 工单数
    n = 0
    for k, v in df_low_vol_fault_current_month.loc[df_low_vol_fault_current_month.一级分类 == '低压故障'].groupby(
            ['单位']).groups.items():
        fault_order[k].append(len(v))
        n += len(v)
    fault_order['合计'].append(n)
    for k in unit_sort:
        fault_order[k].append(round(fault_order[k][6] / fault_order[k][5], 1))
        fault_order[k].append(round(fault_order[k][6] / fault_order[k][0] * 100, 2))
        fault_order[k].append(round(100.0 - fault_order[k][4] - fault_order[k][8], 2))
    mid_vol = sorted(fault_order.items(), key=lambda a: a[1][3], reverse=True)

    low_vol = sorted(fault_order.items(), key=lambda a: a[1][7], reverse=True)
    s_5_1 = "设备与工单关联分析\n（1）故障引发工单情况\n结合本期故障数据与工单数据进行分析。本期抢修工单总数为{}张，其中中压工单{}张，占比{}%，低压工单{}张，占比{}%，其余工单占比{}%。".format(
        len(df_low_vol_fault_current_month), fault_order['合计'][2],
        fault_order['合计'][3], fault_order['合计'][6], fault_order['合计'][8],
        fault_order['合计'][9])
    s_5_1_2 = "中压故障每故障平均工单{}张，{}每故障平均工单数最高，平均每起中压故障引发工单{}张，{}最低，为{}张；低压故障每故障平均工单{}张，{}每故障平均工单数最高,平均每起低压故障引发工单{}张，{}最低，为{}张。".format(
        fault_order['合计'][3], mid_vol[0][0],
        mid_vol[0][1][3], mid_vol[-1][0], mid_vol[-1][1][3], fault_order['合计'][7],
        low_vol[0][0], low_vol[0][1][7], low_vol[-1][0], low_vol[-1][1][7])
    document.add_paragraph(s_5_1)
    document.add_paragraph(s_5_1_2)
    s_5_2 = "（2）抢修效能分析\n本期各类抢修工单办结平均所用时长为{}分钟，按照故障类型计算抢修平均时长见表13；上月抢修工单平均办结时长最长的10个队伍，以及抢修次数最多的10个队伍的故障数和故障平均办结时长见表14、表15。"
    cause_list = ['低压故障', '非电力故障', '计量故障', '客户内部故障', '高压故障']
    cause_dict = {}
    cost_all = 0.0
    for cause in cause_list:
        cause_dict[cause] = []
    for k, v in df_low_vol_fault_current_month.groupby(['一级分类']).groups.items():
        cause_cost = 0.0
        cause_dict[k].append(len(v))  # 到达现场用时（分钟）  工单处理时长（分钟）
        df_tmp = df_low_vol_fault_current_month
        df_tmp = df_tmp.where(df_tmp['关联主单'].isnull()).where(~df_tmp['通知抢修时间'].isnull()).where(
            ~df_tmp['记录到达时间'].isnull()).where(~df_tmp['勘察汇报时间'].isnull()).where(~df_tmp['记录修复时间'].isnull()).where(
            ~df_tmp['恢复送电时间'].isnull()).where(~df_tmp['归档时间'].isnull())
        df_tmp = df_tmp.loc[df_tmp['一级分类'] == k]
        cause_cost += np.sum(df_tmp['工单处理时长（分钟）']) + np.sum(df_tmp['供电所派工用时（分钟）']) + np.sum(
            df_tmp['到达现场用时（分钟）']) + np.sum(df_tmp['派工用时（分钟）'])
        cause_dict[k].append(cause_cost / len(df_tmp))
        cost_all += cause_cost / len(df_tmp)

    s_5_2 = s_5_2.format(round(cost_all / 5, 1))
    document.add_paragraph(s_5_2)

    table = document.add_table(rows=7, cols=3, style='Table Grid')
    table.style.font.size = Pt(12)
    table.cell(0, 0).merge(table.cell(0, 2))
    table.rows[0].cells[0].text = '表13 抢修工单平均办结时长'
    table.rows[1].cells[0].text = '故障类型'
    table.rows[1].cells[1].text = '故障数量'
    table.rows[1].cells[2].text = '平均办结时长（分钟）'
    for row, obj_row in enumerate(table.rows[2:]):
        for col, cell in enumerate(obj_row.cells[0:1]):
            cell.text = cause_list[row]
    for row, obj_row in enumerate(table.rows[2:]):
        for col, cell in enumerate(obj_row.cells[1:]):
            cell.text = str(int(cause_dict[cause_list[row]][col]))

    document.add_paragraph('\n')
    table = document.add_table(rows=10, cols=11, style='Table Grid')
    table.style.font.size = Pt(12)
    table.cell(0, 0).merge(table.cell(0, 10))
    table.rows[0].cells[0].text = '表12 中低压故障与工单关联数据'
    table.cell(1, 0).merge(table.cell(2, 0))
    table.rows[1].cells[0].text = '单位'
    table.cell(1, 1).merge(table.cell(2, 1))
    table.rows[1].cells[1].text = '抢修工单总量'
    table.cell(1, 2).merge(table.cell(1, 5))
    table.rows[1].cells[2].text = '配网中压'
    table.rows[2].cells[2].text = '故障数'
    table.rows[2].cells[3].text = '工单数'
    table.rows[2].cells[4].text = '每故障工单数'
    table.rows[2].cells[5].text = '工单占比(%)'

    table.cell(1, 6).merge(table.cell(1, 9))
    table.rows[1].cells[6].text = '配网低压'
    table.rows[2].cells[6].text = '故障数'
    table.rows[2].cells[7].text = '工单数'
    table.rows[2].cells[8].text = '每故障工单数'
    table.rows[2].cells[9].text = '工单占比(%)'

    table.cell(1, 10).merge(table.cell(2, 10))
    table.rows[2].cells[10].text = '其他工单占比(%)'

    for row, obj_row in enumerate(table.rows[3:]):
        for col, cell in enumerate(obj_row.cells[0:1]):
            cell.text = unit_sort[row]
    for row, obj_row in enumerate(table.rows[3:]):
        for col, cell in enumerate(obj_row.cells[1:]):
            cell.text = str(fault_order[unit_sort[row]][col])
    d_cost = {}
    import re
    for k, v in df_low_vol_fault_current_month.groupby(['抢修队伍/班组']).groups.items():
        count = len(
            df_low_vol_fault_current_month.where(df_low_vol_fault_current_month['抢修队伍/班组'] == k)["工单处理时长（分钟）"].dropna())
        if ',' not in k:
            k_s = re.sub('\(.*?\)', '', k)
            d_cost[k_s] = {}
            d_cost[k_s]['故障数'] = len(df_low_vol_fault_current_month.loc[df_low_vol_fault_current_month['抢修队伍/班组'] == k])
            d_cost[k_s]['故障平均办结时长'] = int(np.sum(df_low_vol_fault_current_month.loc[df_low_vol_fault_current_month['抢修队伍/班组'] == k]['工单处理时长（分钟）']) / count)
    # d_cost
    team_name = sorted(d_cost, key=lambda a: d_cost[a]["故障平均办结时长"], reverse=True)[1:11]
    team_dict = {}
    for idx, team in enumerate(team_name):
        team_dict[idx] = []
        team_dict[idx].append(team)
        team_dict[idx].append(d_cost[team]['故障平均办结时长'])
        team_dict[idx].append(d_cost[team]['故障数'])
    document.add_paragraph('\n')
    table = document.add_table(rows=12, cols=3, style='Table Grid')
    table.style.font.size = Pt(12)
    table.cell(0, 0).merge(table.cell(0, 2))
    table.rows[0].cells[0].text = '表13 抢修工单平均办结时长最长'
    table.rows[1].cells[0].text = '队伍名称'
    table.rows[1].cells[1].text = '故障平均办结时长'
    table.rows[1].cells[2].text = '故障数'
    "table.rows", len(table.rows)
    for row, obj_row in enumerate(table.rows[2:12]):
        for col, cell in enumerate(obj_row.cells):
            cell.text = str(team_dict[row][col])

    ### 处理故障最多
    team_dict = {}
    for k, v in df_low_vol_fault_current_month.groupby(['抢修队伍/班组']).groups.items():
        team_dict[k] = len(v)


    team_fix_dict = {}
    idx = 0
    for t in sorted(team_dict.items(), key=lambda a: a[1], reverse=True)[:11]:
        if "市区抢修指挥班" not in t[0]:
            team_fix_dict[idx] = []
            team_fix_dict[idx].append(re.sub('\(.*?\)', '', t[0]))
            team_fix_dict[idx].append(t[1])
            team_fix_dict[idx].append(round(np.sum(df_low_vol_fault_current_month.loc[df_low_vol_fault_current_month['抢修队伍/班组'] == t[0]]["工单处理时长（分钟）"]) / len(df_low_vol_fault_current_month.where(df_low_vol_fault_current_month['抢修队伍/班组'] == t[0])["工单处理时长（分钟）"].dropna())))
            idx += 1
    document.add_paragraph('\n')
    table = document.add_table(rows=12, cols=3, style='Table Grid')
    table.style.font.size = Pt(12)
    table.cell(0, 0).merge(table.cell(0, 2))
    table.rows[0].cells[0].text = '表14 处理故障最多队伍相关数据'
    table.rows[1].cells[0].text = '队伍名称'
    table.rows[1].cells[2].text = '故障平均办结时长'
    table.rows[1].cells[1].text = '故障数'
    for row, obj_row in enumerate(table.rows[2:12]):
        for col, cell in enumerate(obj_row.cells):
            #         fault_order[unit_sort[row]][col]
            cell.text = str(team_fix_dict[row][col])
    for img in os.listdir(os.getcwd()):
        if 'png' in img:
            os.remove(os.path.join(os.getcwd(),img))
    document.save(os.path.join(os.getcwd(), 'default.docx'))

if __name__ == '__main__':
    gengerate_note()
